#!/usr/bin/env python3
"""One-off generator for MoodRefreshment_Technical_Inventory.docx. Run: python3 scripts/generate-tech-inventory-docx.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "MoodRefreshment_Technical_Inventory.docx"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for c, text in enumerate(row):
            cells[c].text = text
    doc.add_paragraph()


def main():
    doc = Document()
    t = doc.add_heading("Mood Refreshment App — Technical inventory", 0)
    t.runs[0].font.size = Pt(18)

    doc.add_paragraph(
        "Document for defence / readers: React Native + Expo stack, npm dependencies, "
        "project modules, screens, and where each part is used. Generated for the mobile app under mobile/."
    )

    doc.add_heading("1. Architecture (big pieces)", level=1)
    add_table(
        doc,
        ["Piece", "What it is", "Where in the project"],
        [
            ["React", "Components, state, hooks", "All pages, components, MoodContext"],
            ["React Native", "Native UI (View, Text, Animated, …)", "Screens and components"],
            ["Expo", "SDK, CLI, native modules", "App.js, app.json, expo-* packages"],
            [
                "React Navigation",
                "Stack navigation",
                "mobile/App.js — NavigationContainer, NativeStackNavigator",
            ],
            [
                "MoodContext",
                "Global user, moods, journals, settings",
                "mobile/src/context/MoodContext.js; pages use useContext",
            ],
            [
                "REST API",
                "HTTP to Node server",
                "axios / api in MoodContext, Login, ChatRoom, Profile, admin pages, Motivation, …",
            ],
            [
                "Socket.IO client",
                "Realtime chat",
                "mobile/src/pages/EmotionRooms.js, ChatRoom.js — io(SOCKET_URL)",
            ],
            [
                "AsyncStorage",
                "On-device key–value store",
                "MoodContext.js, reportClientId.js",
            ],
            ["Theme", "Colours & typography", "mobile/src/styles/theme.js"],
        ],
    )

    doc.add_heading("2. package.json dependencies — used in source", level=1)
    add_table(
        doc,
        ["Dependency", "Role", "Main locations"],
        [
            ["react, react-native", "Core UI", "Entire mobile app"],
            ["expo", "Platform / toolchain", "App.js, config, Expo modules"],
            ["@expo/metro-runtime", "Bundler runtime", "Loaded by Expo (no direct import)"],
            [
                "react-dom, react-native-web",
                "Web target",
                "Used when running Expo web (app.json lists web platform)",
            ],
            [
                "@react-navigation/native, native-stack",
                "Navigation",
                "App.js",
            ],
            [
                "react-native-screens, react-native-safe-area-context",
                "Native stack & safe areas",
                "App.js SafeAreaProvider; SafeScreen.js useSafeAreaInsets",
            ],
            ["axios", "HTTP client", "MoodContext, pages, api.js"],
            ["socket.io-client", "Chat websocket", "EmotionRooms.js, ChatRoom.js"],
            [
                "@react-native-async-storage/async-storage",
                "Persistence",
                "MoodContext.js, reportClientId.js",
            ],
            [
                "@react-native-community/slider",
                "Slider UI",
                "MoodCheck.js",
            ],
            [
                "@react-native-google-signin/google-signin",
                "Google Sign-In",
                "Login.js",
            ],
            [
                "expo-font + @expo-google-fonts/*",
                "Custom fonts",
                "App.js useFonts",
            ],
            ["expo-status-bar", "Status bar style", "App.js"],
            ["expo-constants", "Config / manifest values", "config.js, Login.js"],
            [
                "expo-linear-gradient",
                "Gradients",
                "Splash.js, Motivation.js, AmbientBackground.js, Button.js",
            ],
            ["expo-blur", "Blur views (iOS cards)", "Card.js"],
            ["expo-haptics", "Touch feedback", "Button.js"],
            ["expo-notifications", "Local scheduling", "Profile.js (dynamic require)"],
            ["lucide-react-native", "Icons", "Login.js"],
            ["react-native-svg", "SVG", "Home.js, Breathing.js"],
            [
                "@expo/vector-icons (Feather)",
                "Icons",
                "Many pages",
            ],
        ],
    )

    doc.add_heading("3. Declared in package.json but not imported in .js (current codebase)", level=1)
    doc.add_paragraph(
        "These may be leftovers or reserved for future use; they do not appear in mobile/src imports today."
    )
    add_table(
        doc,
        ["Package", "Note"],
        [
            [
                "styled-components",
                "Listed in package.json; no styled/ import under mobile/src.",
            ],
            [
                "expo-crypto",
                "Listed in package.json; no import found under mobile/src.",
            ],
        ],
    )

    doc.add_heading("4. Config / build tooling", level=1)
    add_table(
        doc,
        ["Item", "Role"],
        [
            [
                "expo-web-browser (app.json plugin)",
                "Native wiring for in-app browser when building; useful for some OAuth flows.",
            ],
            [
                "expo-dev-client",
                "Custom development builds with native modules beyond Expo Go.",
            ],
        ],
    )

    doc.add_heading("5. Project modules (mobile/src/utils)", level=1)
    add_table(
        doc,
        ["File", "Purpose", "Used in"],
        [
            ["config.js", "API_URL, SOCKET_URL, Google IDs, dev host", "API/socket callers"],
            ["api.js", "Axios instance + fallback URLs", "Login.js"],
            ["debug.js", "testConnection", "Login.js"],
            ["deviceTime.js", "Device timezone / locale dates", "ChatRoom, Journal, reports, AdminDashboard, moodAnalytics"],
            ["moodAnalytics.js", "Streak, weekly, badges, distress helper", "MoodContext, MoodCheck, MoodInsights"],
            ["reportClientId.js", "Stable guest key for reports", "ChatRoom.js, MyChatReports.js"],
            ["identity.js", "Display identity helpers", "Profile.js"],
            ["analytics.js", "Lightweight tracking hook", "MoodCheck.js"],
        ],
    )

    doc.add_heading("6. Components (mobile/src/components)", level=1)
    add_table(
        doc,
        ["Component", "Role", "Used in"],
        [
            ["SafeScreen", "Safe area + optional ambient background", "Most full screens"],
            ["AmbientBackground", "Animated gradients", "Via SafeScreen"],
            ["Card", "Card; BlurView on iOS", "Many pages"],
            ["Button", "Primary button, haptics, gradient", "Login, MoodCheck, Journal, Profile, …"],
            ["Input", "Text field styling", "Forms where imported"],
            ["FadeInView", "Entrance animation", "Home.js"],
            ["ErrorBoundary", "Catch render errors", "App.js"],
        ],
    )

    doc.add_heading("7. Screens — main tech per page", level=1)
    add_table(
        doc,
        ["Screen", "Notable dependencies / behaviour"],
        [
            ["Splash", "expo-linear-gradient, Animated"],
            ["Login", "GoogleSignin, api, lucide, Constants, testConnection"],
            ["Home", "react-native-svg, FadeInView, MoodContext"],
            ["MoodCheck", "Slider, moodAnalytics, analytics"],
            ["MoodInsights", "MoodContext, moodAnalytics"],
            ["Journal / JournalEntry", "MoodContext, deviceTime, Card, Button"],
            ["Breathing", "react-native-svg, Animated"],
            ["EmotionRooms", "socket.io-client, room list"],
            ["ChatRoom", "socket.io-client, axios, reportClientId, deviceTime"],
            ["Motivation", "axios, LinearGradient, daily drops API"],
            ["Profile", "axios, expo-notifications, identity, settings"],
            ["AdminDashboard", "axios, useFocusEffect, BackHandler, moods/users"],
            ["AdminChatReports / MyChatReports", "axios, chat-reports API"],
        ],
    )

    doc.add_heading("8. Async / AsyncStorage (quick reference)", level=1)
    doc.add_paragraph(
        "async/await: used to wait on network (axios) and storage without blocking the UI — see MoodContext, Login, "
        "ChatRoom, JournalEntry, Profile, admin pages, utils/reportClientId, utils/debug."
    )
    doc.add_paragraph(
        "AsyncStorage (library name): key–value persistence on the phone — not the same keyword as async/await. "
        "Used in MoodContext (user, settings, caches), reportClientId (guest report key)."
    )

    doc.add_heading("9. One-sentence summary (oral defence)", level=1)
    doc.add_paragraph(
        "The app uses React Native with Expo: React Navigation for screens, MoodContext and AsyncStorage for state "
        "and persistence, axios for REST and socket.io-client for chat, plus Expo modules for fonts, gradients, blur, "
        "haptics, notifications, and Google Sign-In; src/utils holds date handling, analytics, reporting id, and API fallbacks."
    )

    doc.add_paragraph("")
    doc.add_paragraph("— End of document —")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
